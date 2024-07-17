from flask import Flask, request, jsonify
from docx import Document
from docx2pdf import convert
from flask_cors import CORS
import os
import uuid
import pythoncom
from BtReport import BtReport

app = Flask(__name__)
CORS(app)  # CORS define

def checkText(text):
    print("normal beyin bt" in text.lower())
    return "normal beyin bt" in text.lower()

normalBt = "Bilateral mastoid sellüllerin aerasyonları normaldir. Her iki bulbus oculi, optik sinir, orbital yapılar normal görünümdedir. Kranyoservikal bileşke normal olarak izlenmektedir. lV.ventrikül ve bazal sisternalar normal genişliktedir. Serebellar folialar doğal görünümdedir. Beyin sapı, pons ve bulbus sinyal intensiteleri ve morfolojilerinde anlamlı bir değişiklik saptanmamıştır. lll.ventrikül ve lateral ventriküller normal genişliktedir. Serebral fissür ve sulkusların derinlik ve genişlikleri hasta yaşı ile uyumludur. Bilateral talamus ve bazal ganglionlar normal görünümdedir. Korpus kallozum kalınlığı ve sinyal intensitesi normal olarak izlenmektedir. Gri ve ak madde sinyal intensitelerinde anlamlı bir değişiklik izlenmemiştir. Sonuç ve Öneriler: Normal sınırlarda beyin MR bulguları."

def json_to_pdf(json_file, pdf_file):
    # Initialize COM library
    pythoncom.CoInitialize()

    try:
        # If PDF file exists, remove it
        if os.path.exists(pdf_file):
            os.remove(pdf_file)

        # If DOCX file exists, remove it
        docx_file = pdf_file[:-3] + 'docx'
        if os.path.exists(docx_file):
            os.remove(docx_file)

        text = json_file.get('reportText', '')

        title = "Teknik: FSE T2 axial, sagittal, FSE T1 axial, Flair T2 koronal"
        subtitle = "İnceleme: KRANİYAL MR Bulgular."

        sentences = normalBt.split(". ")
        report = BtReport(title, subtitle, *sentences)

        doc = Document()

        if checkText(text):
            doc.add_paragraph(title)
            doc.add_paragraph(subtitle)
            for sentence in normalBt.split(". "):
                doc.add_paragraph(sentence)
        else:
            new_sentences = text.split(". ")
            for sentence in new_sentences:
                print(sentence)
                if "serebral fissür" in sentence.lower():
                    report.sentences[7] = sentence

                if "t2" in sentence.lower():
                    report.sentences[7] = "Her iki parietal ve oksipital lobda beyaz cevher yerleşimli T2A FLAIRda non spesifik milimetrik gliotik odaklar izlendi. "
                    report.sentences[8] = "Paranazal sinüslerde mukozal kalınlaşma mevcuttur. "
                if "beyaz cevherde" in sentence.lower():
                    report.sentences[9] = sentence
                if "temporal" in sentence.lower():
                    report.sentences[6] = sentence

            report.sentences[11] = "Sonuç ve Öneriler: Mevcut bulgular rapor içerisinde belirtilmiştir."

            doc.add_paragraph(str(report))

        doc.save(docx_file)
        convert(docx_file, pdf_file)

    finally:
        pythoncom.CoUninitialize()

    return file_path_to_url(pdf_file)

def file_path_to_url(file_path):

    file_path = file_path.replace(" ", "%20")
    url = "http://192.168.56.1:8080/" + file_path
    print(url)
    return url

@app.route('/convert', methods=['POST'])
def convert_endpoint():
    json_data = request.get_json()
    pdf_file = str(uuid.uuid4()) + '.pdf'

    pdf_url = json_to_pdf(json_data, pdf_file)
    return jsonify({'text': pdf_url})

if __name__ == '__main__':
    print("App started")
    app.run(debug=True, port=5001)



# npx http-server ./
