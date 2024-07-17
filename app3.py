from flask import Flask, request, jsonify
from docx import Document
from docx2pdf import convert
from flask_cors import CORS
import os
import uuid

app = Flask(__name__)
CORS(app) 

@app.route('/convert', methods=['POST'])
def json_to_pdf():
    json_data = request.get_json()
    text = json_data.get('reportText', '')

    title = "Teknik:FSE T2 axial,sagittal,FSE T1 axial,Flair T2 koronal"
    subtitle = "İnceleme:KRANİYAL MR Bulgular."

    pdf_file = str(uuid.uuid4()) + '.pdf'

    if not os.path.exists(pdf_file):
        doc = Document()
        doc.add_paragraph(title)
        doc.add_paragraph(subtitle)
        doc.add_paragraph(text)
        doc.save(pdf_file[:-3] + 'docx')
        #BURAYI YORUMLA DENEME

        convert(pdf_file[:-3] + 'docx', pdf_file)
    else:
        docx_file = pdf_file[:-3] + 'docx'
        convert(docx_file, pdf_file)

    print(pdf_file)
    pdf_url = file_path_to_url(pdf_file)
    return jsonify({'text': pdf_url})

def file_path_to_url(file_path):
    
    file_path = file_path.replace(" ", "%20")
    url = "http://192.168.56.1:8080/" + file_path
    print(url)
    return url

if __name__ == '__main__':
    print("App started")
    app.run(debug=True, port=5001)


#SERVER : npx http-server ./
